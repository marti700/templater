from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# document = Document("EjemploGenerales.docx")
# paragraph = document.paragraphs[0]
# print(paragraph.text)

# document = Document("EjemploGenerales.docx")
# with open('resume.xml', 'w') as f:
# 	f.write(document._element.xml)


# document = Document()
# document.add_heading('Document Title', 0)

# p = document.add_paragraph('Entre de una parte El senor ')
# p.add_run('Teodoro Alcantara Aquino ').bold = True
# p.add_run('Dominicano, mayor de edad, soltero, empleado privado, portador de la cedula de identidad y electoral No. 40221849o15, domiciliado y recidente en Santo Domingo;')
# p.add_run('quien en lo que sigue del presente contrato se denominara ')
# p.add_run('EL VENDEDOR').bold=True
# p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# document.save('demo.docx')


import re

sections = ["<<vendedor>>,", "<<comprador>>",
            "<<justificacion>>", "<<descripcion>>"]
data_json_simaltion = {
    "<<vendedor>>": [
        {
            "nombre": "Teodoro ",
            "apellido": "Alcantara Aquino",
            "ocupacion": "ingeniero",
            "nacionalidad": "Dominicana",
            "direccion": "C/ Antonio Maceo, edif Coronado, apto 5A, Santo Domingo, Distrino Nacional",
            "identificacion": "cedula",
            "numero_identificacion": "40221884915",
            "estado_civil": "soltero"
        },
        {
            "nombre": "Loanny",
            "apellido": "Alcantara Mojica",
            "ocupacion": "Doctora",
            "nacionalidad": "Dominicana",
            "direccion": "C/ Antonio Maceo, edif Coronado, apto 5A, Santo Domingo, Distrino Nacional",
            "identificacion": "cedula",
            "numero_identificacion": "40221884915",
            "estado_civil": "soltera"
        }
    ],
    "<<comprador>>": [
        {
            "nombre": "Anulfo",
            "apellido": "Alcantara Aquino",
            "ocupacion": "Profesor",
            "nacionalidad": "Dominicana",
            "direccion": "C/ Olmedo Paniagua, #20 Res olmedo paniagua, San Juan de la Maguana",
            "identificacion": "cedula",
            "numero_identificacion": "40251948812",
            "estado_civil": "soltero"
        }
    ],
    "<<descripcion>>": ["Un carro"],
    "<<justificacion>>": ["Es mio"]
}

def create_text_object(template):
    template2 = "{nombre:negrita,mayusculas}, de nacionalidad {nacionalidad}, mayor de edad, {estado_civil}, {ocupacion}, portador de {identificacion} No. {numero_identificacion}, domiciliado y residente en {direccion}"
    text_object = {
        "paragraphs": []
    }

    # Split the template into words and punctuation
    words = re.split(r"\s+", template)

    # Iterate over the words and create text nodes
    for word in words:
        match = re.match(r"{(.+):(.+)}", word)
        # match2 = re.match(r"{(.+)}", word)
        if match:
            # Found a curly-braced expression
            string = match.group(1)
            styles = match.group(2).split(",")
            text_object["paragraphs"].append({
                "textNode": {
                    "string": "{"+string+'}',
                    "style": styles
                }
            })
        # elif match2:
        #     # Found a curly-braced expression
        #     string = match2.group(1)
        #     text_object["paragraphs"].append({
        #         "textNode": {
        #             "string": word,
        #             "style": []
        #         }
        #     })
        else:
            # Regular word or punctuation
            text_object["paragraphs"].append({
                "textNode": {
                    "string": word,
                    "style": []
                }
            })

    return text_object


def enhanceTemplate(main_tempate):

    template = "{nombre:negrita,mayusculas}, de nacionalidad {nacionalidad}, mayor de edad, {estado_civil}, {ocupacion}, portador de {identificacion} No. {numero_identificacion}, domiciliado y residente en {direccion}"
    templates = {
        "<<vendedor>>": template,
        "<<comprador>>": template,
        "<<descripcion>>": "Un carro",
        "<<justificacion>>": "Es mio"
    }

    for key, value in data_json_simaltion.items():
        res = ["" + templates[key]
               for r in range(len(value)) if key in templates]
        if key in templates:
            main_tempate = main_tempate.replace(key, "; ".join(res))
    return main_tempate


def buildDocument(params, template):
    # document = Document("EjemploGenerales.docx")
    # paragraph = document.paragraphs[0]
    # print(paragraph.text)

    # document = Document("EjemploGenerales.docx")
    # with open('resume.xml', 'w') as f:
    # 	f.write(document._element.xml)

    document = Document()
    document.add_heading('Document Title', 0)

    buffer = ''
    p = document.add_paragraph(buffer)
    for word in template['paragraphs']:
        w: str = ''
        w = word['textNode']['string']
        if w[0] == '{':
            p.add_run(buffer + ' ')
            replace_text_in_paragraph(word, params, p)
            buffer = ''
            # params.pop(0)
        else:
            buffer = buffer + w + ' '

    # p = document.add_paragraph('Entre de una parte El senor ')
    # p.add_run('Dominicano, mayor de edad, soltero, empleado privado, portador de la cedula de identidad y electoral No. 40221849o15, domiciliado y recidente en Santo Domingo;')
    # p.add_run('quien en lo que sigue del presente contrato se denominara ')
    # p.add_run('EL VENDEDOR').bold=True
    # p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.save('demo1.docx')


def replace_text_in_paragraph(word, params, paragraph):
    w = word['textNode']['string']
    w = w[:-1]
    w = w.replace('{', '')
    w = w.replace('}', '')
    text = ''
    if w[0] != '[':
        # TODO search the json
        text = params[0].pop(w, None) if len(params) > 0 else "Hay bobo"
        if text == None:
            params.pop(0)
            text = params[0].pop(w, None)
    else:
        text = w[1:-1]
    # fromJson = 'TEST_TEXT'
    if len(word['textNode']['style']) > 0:
        styles = word['textNode']['style']
        run = paragraph.add_run(text + ' ')
        for s in styles:
            if s == 'mayusculas':
                run.text = text.upper() + ' '
            if s == 'negrita':
                run.bold = True
    else:
        paragraph.add_run(text + '')

def template_to_paragraphs(template):
    return re.split("\n",template)

# Example usage
template2 = "Entre de una parte <<vendedor>>, quien en lo que sigue del presente contrato se denominará {[vendedor]:negrita,mayusculas}, y de la otra parte: <<comprador>>, quien en lo que sigue del presente contrato se denominará {[comprador]:negrita,mayusculas}, se ha convenido y pactado el siguiente"

template2 = enhanceTemplate(template2)

text_object = create_text_object(template2)
print(text_object)

params = []
for key, value in data_json_simaltion.items():
    params.append(value)

params = [item for values in params for item in values]
buildDocument(params, text_object)

p = template_to_paragraphs("""Entre de una parte <<vendedor>>, quien en lo que sigue del presente contrato se denominará {rol1:negrita:mayusculas}, y de la otra parte: <<comprador>>, quien en lo que sigue del presente contrato se denominará {rol2:negrita:mayusculas, se ha convenido y pactado el siguiente: -----------------------------------------------------
Contrato
PRIMERO: {rol1:negrita:mayusculas}, por medio del presente acto VENDE, CEDE Y TRANSFIERE desde ahora y para siempre con todas las garantías legales y sin impedimento alguno a favor de {rol2:negrita:mayusculas}, quien acepta conforme el bien  de su propiedad que se describe a continuación:

<<descripcion>>

SEGUNDO: El precio convenido y pactado entre las partes ha sido por la suma de {precio_letras}(RD$ {precio_numero}), pesos Dominicanos, suma esta que {rol1:negrita:mayusculas}, declara haber recibido en moneda de curso legal de manos de {rol2:negrita:mayusculas} a su entera satisfacción, por lo que otorga  recibo de descargo y finiquito legal y carta liberatoria de pago en toda forma de derecho a favor de {rol2:negrita:mayusculas}.-
======================================================
TERCERO: {rol1:negrita:mayusculas}, justifica su derecho de propiedad al objeto de la presente venta por <<justificacion>>.

CUARTO: Las partes hacen constar que están conformes con todos los convenidos y pactado en el presente contrato, por lo que proceden a dar su aprobación firmando al pie del mismo.---------------------------------------------

HECHO LEIDO Y FIRMADO DE BUENA FE, en dos  (2) originales del mismo contenido y efecto uno para cada una de las partes. En  la ciudad y municipio de San Juan de la Maguana, provincia San Juan, a los {fecha}.


                   Todo bien                                                            asdlfj
                {rol1:negrita}                                                    {rol2:negrita}


YO DR. TEODORO ALCANTARA BIDO, Notario Público de los del número de este municipio de San Juan de la Maguana, miembro del colegio dominicano de Notarios, con matrícula No. 1799, CERTIFICO Y DOY FE, que las firmas que aparecen al pie del presente documento fueron puestas libre y voluntariamente en mi presencia, por los Señores: Todo bien Y asdlfj, de generales anotadas y quienes me declararon BAJO LA FE DEL JURAMENTO, que esas son las mismas firmas que ellos acostumbran a usar en todos los actos de  su vida pública y privada, por lo que merecen Fe y Crédito. En  la ciudad y municipio de San Juan de la Maguana, provincia San Juan, a los {fecha}.---------------------------------------------------


                                                                 DR. TEODORO ALCÁNTARA BIDO
                                                                                          Notario-Publico
""")
print("OK")